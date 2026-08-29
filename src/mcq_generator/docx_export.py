"""
docx_export.py
---------------
Xuất bộ câu hỏi MCQ ra file Word (.docx), trình bày theo đúng cấu trúc của
file mẫu người dùng cung cấp ("Sheet 1, dòng 2"):

    Câu N. <mã năng lực> -- <tên năng lực>
    Mức độ tư duy: <...>
    Phần thân
    <đoạn văn tình huống>
    <câu hỏi dẫn in đậm>
    A. ...
    B. ...
    C. ...
    D. ...
    Đáp án: <X>.
    Kiểm tra kỹ thuật: <ghi chú/cảnh báo từ validator, nếu có>

    ... (lặp lại cho từng câu) ...

    Bảng tóm tắt cấu trúc năng lực ở cuối tài liệu.
"""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .validator import ValidationResult


def _add_heading(doc: Document, text: str, size: int = 13, bold: bool = True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def export_to_docx(
    questions: list[dict],
    output_path: str,
    scenario_title: str = "Tình huống lâm sàng",
    raw_scenario: str | None = None,
    validation_results: list[ValidationResult] | None = None,
) -> str:
    """Ghi bộ câu hỏi ra file .docx. Trả về output_path để tiện dùng nối chuỗi."""
    doc = Document()

    title = doc.add_heading(f"BỘ {len(questions)} CÂU HỎI MCQ THEO NĂNG LỰC", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub_run = sub.add_run(scenario_title)
    sub_run.italic = True
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if raw_scenario:
        doc.add_paragraph()
        _add_heading(doc, "Nguồn tình huống thô:", size=11)
        doc.add_paragraph(raw_scenario.strip())

    validation_by_number = {
        v.question_number: v for v in (validation_results or [])
    }

    for q in questions:
        doc.add_paragraph()  # dòng trống ngăn cách
        n = q.get("question_number", "?")
        code = q.get("competency_code", "")
        name = q.get("competency_name", "")
        level = q.get("thinking_level", "")

        _add_heading(doc, f"Câu {n}. {code} -- {name}")
        level_p = doc.add_paragraph()
        level_run = level_p.add_run(f"Mức độ tư duy: {level}")
        level_run.italic = True

        _add_heading(doc, "Phần thân", size=11)
        doc.add_paragraph(q.get("stem", ""))

        lead_in_p = doc.add_paragraph()
        lead_in_run = lead_in_p.add_run(q.get("lead_in", ""))
        lead_in_run.bold = True

        options = q.get("options", {})
        for letter in ["A", "B", "C", "D"]:
            opt_p = doc.add_paragraph()
            opt_run = opt_p.add_run(f"{letter}. ")
            opt_run.bold = True
            opt_p.add_run(str(options.get(letter, "")))

        answer_p = doc.add_paragraph()
        answer_run = answer_p.add_run(f"Đáp án: {q.get('correct_answer', '')}.")
        answer_run.bold = True

        rationale = q.get("rationale")
        if rationale:
            rationale_p = doc.add_paragraph()
            rationale_label = rationale_p.add_run("Giải thích: ")
            rationale_label.bold = True
            rationale_p.add_run(rationale)

        tech_p = doc.add_paragraph()
        tech_label = tech_p.add_run("Kiểm tra kỹ thuật: ")
        tech_label.bold = True

        v = validation_by_number.get(n)
        if v and v.warnings:
            tech_p.add_run(
                "Có " + str(len(v.warnings)) + " điểm cần xem lại — "
                "" + "; ".join(v.warnings)
            )
        else:
            tech_p.add_run("Không phát hiện lỗi kỹ thuật tự động (heuristic).")

    # Bảng tóm tắt cuối tài liệu — giống bảng "Tóm tắt cấu trúc năng lực"
    # trong file mẫu.
    doc.add_paragraph()
    doc.add_heading("Tóm tắt cấu trúc năng lực", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
        "Câu", "Mã năng lực -- Tên", "Mức độ tư duy", "Đáp án",
    )
    for q in questions:
        row = table.add_row().cells
        row[0].text = str(q.get("question_number", ""))
        row[1].text = f"{q.get('competency_code', '')} -- {q.get('competency_name', '')}"
        row[2].text = str(q.get("thinking_level", ""))
        row[3].text = str(q.get("correct_answer", ""))

    doc.save(output_path)
    return output_path
